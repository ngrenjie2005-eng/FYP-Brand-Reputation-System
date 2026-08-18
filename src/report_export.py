# ============================================================
# BRANDPULSE AI
# PROFESSIONAL REPORT EXPORT
#
# Supported formats:
# 1. Microsoft Word (.docx)
# 2. PDF (.pdf)
#
# Predictive AI:
# - DistilBERT
#
# Department LLMs:
# - OpenRouter
# - Ollama
#
# Executive LLM:
# - Gemini
# ============================================================

from datetime import datetime
from io import BytesIO
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from fpdf import FPDF


# ============================================================
# GENERAL HELPERS
# ============================================================

def current_timestamp():
    """
    Return the current report generation
    date and time in a readable format.
    """

    return datetime.now().strftime(
        "%d %B %Y, %H:%M"
    )


def clean_markdown_text(
    text,
):
    """
    Remove simple Markdown formatting
    from generated LLM content.
    """

    if text is None:
        return ""

    text = str(text)

    replacements = {
        "**": "",
        "__": "",
        "`": "",
    }

    for (
        original,
        replacement,
    ) in replacements.items():

        text = text.replace(
            original,
            replacement,
        )

    return text.strip()


# ============================================================
# PDF SAFE TEXT
# ============================================================

def pdf_safe_text(
    text,
):
    """
    Convert common Unicode characters and
    emojis into PDF-safe text.

    Helvetica built into FPDF does not support
    every Unicode character, so unsupported
    characters are simplified or removed.
    """

    if text is None:
        return ""

    text = str(text)


    replacements = {

        # Quotes
        "’": "'",
        "‘": "'",
        "“": '"',
        "”": '"',

        # Dashes
        "–": "-",
        "—": "-",

        # Other punctuation
        "…": "...",
        "•": "-",
        "→": "->",

        # Mathematical
        "≥": ">=",
        "≤": "<=",

        # Status
        "✓": "[OK]",
        "✅": "[OK]",
        "⚠️": "[Warning]",
        "⚠": "[Warning]",

        # Sentiment
        "🟢": "[Positive]",
        "🔴": "[Negative]",
        "🟡": "[Mixed]",

        # AI / visualisation emojis
        "📊": "",
        "📈": "",
        "📉": "",
        "🧠": "",
        "🤖": "",
        "👔": "",
        "🛠️": "",
        "🛠": "",
        "🧩": "",
        "🎧": "",
        "📣": "",
        "💳": "",
        "💡": "",
        "🔎": "",
        "📑": "",
        "📘": "",
        "📕": "",
        "✨": "",
        "🌐": "",
        "🚨": "",
        "🎯": "",
        "📚": "",
        "📏": "",
        "🔗": "",
        "🦙": "",
        "💚": "",
        "💗": "",
        "✍️": "",
        "🚀": "",
    }


    for (
        original,
        replacement,
    ) in replacements.items():

        text = text.replace(
            original,
            replacement,
        )


    text = clean_markdown_text(
        text
    )


    # Replace characters unsupported by
    # built-in Helvetica.
    text = (
        text
        .encode(
            "latin-1",
            errors="replace",
        )
        .decode(
            "latin-1",
        )
    )


    return text


# ============================================================
# DATA HELPERS
# ============================================================

def get_word_and_count(
    item,
):
    """
    Support word-frequency data stored as either:

    {"word": "music", "count": 10}

    or

    ("music", 10)
    """

    if isinstance(
        item,
        dict,
    ):

        return (
            item.get(
                "word",
                "",
            ),
            item.get(
                "count",
                0,
            ),
        )


    if isinstance(
        item,
        (
            list,
            tuple,
        ),
    ):

        word = (
            item[0]
            if len(item) > 0
            else ""
        )

        count = (
            item[1]
            if len(item) > 1
            else 0
        )

        return (
            word,
            count,
        )


    return (
        str(item),
        "",
    )


# ============================================================
# DOCX HELPERS
# ============================================================

def docx_add_title(
    document,
    text,
):

    paragraph = (
        document.add_paragraph()
    )

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    run = paragraph.add_run(
        text
    )

    run.bold = True

    run.font.size = Pt(
        23
    )


def docx_add_centered_text(
    document,
    text,
    size=11,
    bold=False,
):

    paragraph = (
        document.add_paragraph()
    )

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    run = paragraph.add_run(
        str(text)
    )

    run.bold = bold

    run.font.size = Pt(
        size
    )


def docx_add_llm_report(
    document,
    report_text,
):
    """
    Convert basic Markdown-style LLM output
    into Word headings, lists and paragraphs.
    """

    if not report_text:
        return


    lines = (
        str(report_text)
        .splitlines()
    )


    for raw_line in lines:

        line = (
            raw_line.strip()
        )


        if not line:
            continue


        # ----------------------------------------------------
        # H1
        # ----------------------------------------------------

        if line.startswith(
            "# "
        ):

            document.add_heading(
                clean_markdown_text(
                    line[2:]
                ),
                level=1,
            )


        # ----------------------------------------------------
        # H2
        # ----------------------------------------------------

        elif line.startswith(
            "## "
        ):

            document.add_heading(
                clean_markdown_text(
                    line[3:]
                ),
                level=2,
            )


        # ----------------------------------------------------
        # H3
        # ----------------------------------------------------

        elif line.startswith(
            "### "
        ):

            document.add_heading(
                clean_markdown_text(
                    line[4:]
                ),
                level=3,
            )


        # ----------------------------------------------------
        # BULLET
        # ----------------------------------------------------

        elif (
            line.startswith("- ")
            or
            line.startswith("• ")
        ):

            document.add_paragraph(
                clean_markdown_text(
                    line[2:]
                ),
                style="List Bullet",
            )


        # ----------------------------------------------------
        # NUMBERED LIST
        # ----------------------------------------------------

        elif re.match(
            r"^\d+\.\s",
            line,
        ):

            cleaned_line = re.sub(
                r"^\d+\.\s*",
                "",
                line,
            )


            document.add_paragraph(
                clean_markdown_text(
                    cleaned_line
                ),
                style="List Number",
            )


        # ----------------------------------------------------
        # NORMAL PARAGRAPH
        # ----------------------------------------------------

        else:

            document.add_paragraph(
                clean_markdown_text(
                    line
                )
            )


# ============================================================
# CREATE DOCX REPORT
# ============================================================

def create_brandpulse_docx(
    analysis_summary,
    manager_reports,
    executive_report,
):
    """
    Create a complete BrandPulse AI
    Microsoft Word report.

    Returns:
        BytesIO
    """

    document = Document()


    # ========================================================
    # PAGE MARGINS
    # ========================================================

    for section in (
        document.sections
    ):

        section.top_margin = (
            Inches(0.65)
        )

        section.bottom_margin = (
            Inches(0.65)
        )

        section.left_margin = (
            Inches(0.75)
        )

        section.right_margin = (
            Inches(0.75)
        )


    # ========================================================
    # DEFAULT FONT
    # ========================================================

    normal_style = (
        document.styles[
            "Normal"
        ]
    )

    normal_style.font.name = (
        "Arial"
    )

    normal_style.font.size = Pt(
        10.5
    )


    # ========================================================
    # COVER PAGE
    # ========================================================

    docx_add_title(
        document,
        "BrandPulse AI",
    )


    docx_add_centered_text(
        document,
        (
            "AI-Assisted Brand Reputation "
            "Intelligence Report"
        ),
        size=14,
        bold=True,
    )


    document.add_paragraph()


    docx_add_centered_text(
        document,
        (
            "Online Review-Based Brand "
            "Reputation Prediction Using "
            "NLP Techniques"
        ),
        size=12,
    )


    document.add_paragraph()


    docx_add_centered_text(
        document,
        (
            "Predictive Model: DistilBERT"
        ),
        size=10,
    )


    docx_add_centered_text(
        document,
        (
            "Department Generative AI: "
            "OpenRouter and Ollama"
        ),
        size=10,
    )


    docx_add_centered_text(
        document,
        (
            "Executive Generative AI: Gemini"
        ),
        size=10,
    )


    document.add_paragraph()


    docx_add_centered_text(
        document,
        (
            "Generated: "
            + current_timestamp()
        ),
        size=9,
    )


    document.add_page_break()


    # ========================================================
    # 1. BRAND REPUTATION OVERVIEW
    # ========================================================

    document.add_heading(
        (
            "1. Brand Reputation Overview"
        ),
        level=1,
    )


    overview_table = (
        document.add_table(
            rows=1,
            cols=2,
        )
    )


    overview_table.style = (
        "Table Grid"
    )


    overview_table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )


    header = (
        overview_table.rows[
            0
        ].cells
    )


    header[0].text = (
        "Indicator"
    )


    header[1].text = (
        "Result"
    )


    overview_values = [

        (
            "Reviews Analysed",
            analysis_summary.get(
                "total_reviews",
                0,
            ),
        ),

        (
            "Positive Reviews",
            analysis_summary.get(
                "positive_reviews",
                0,
            ),
        ),

        (
            "Negative Reviews",
            analysis_summary.get(
                "negative_reviews",
                0,
            ),
        ),

        (
            "Positive Percentage",
            (
                f"{analysis_summary.get(
                    'positive_percentage',
                    0
                )}%"
            ),
        ),

        (
            "Negative Percentage",
            (
                f"{analysis_summary.get(
                    'negative_percentage',
                    0
                )}%"
            ),
        ),

        (
            "Brand Reputation Score",
            (
                f"{analysis_summary.get(
                    'reputation_score',
                    0
                )}%"
            ),
        ),
    ]


    for (
        indicator,
        result,
    ) in overview_values:

        row = (
            overview_table
            .add_row()
            .cells
        )


        row[0].text = str(
            indicator
        )


        row[1].text = str(
            result
        )


    document.add_paragraph()


    methodology_note = (
        document.add_paragraph()
    )


    methodology_run = (
        methodology_note.add_run(
            "Methodology Note: "
        )
    )


    methodology_run.bold = True


    methodology_note.add_run(
        (
            "The Brand Reputation Score is a "
            "project-defined decision-support "
            "indicator calculated from the "
            "proportion of positive DistilBERT "
            "sentiment predictions. It is not "
            "an official Spotify metric or a "
            "universal industry-standard brand "
            "reputation measurement."
        )
    )


    # ========================================================
    # 2. ISSUE ANALYSIS
    # ========================================================

    document.add_heading(
        (
            "2. Negative Review "
            "Issue Analysis"
        ),
        level=1,
    )


    issue_counts = (
        analysis_summary.get(
            "issue_counts",
            {},
        )
    )


    if issue_counts:

        issue_table = (
            document.add_table(
                rows=1,
                cols=2,
            )
        )


        issue_table.style = (
            "Table Grid"
        )


        issue_table.alignment = (
            WD_TABLE_ALIGNMENT.CENTER
        )


        issue_header = (
            issue_table.rows[
                0
            ].cells
        )


        issue_header[0].text = (
            "Issue Category"
        )


        issue_header[1].text = (
            "Mentions"
        )


        for (
            issue,
            count,
        ) in issue_counts.items():

            row = (
                issue_table
                .add_row()
                .cells
            )


            row[0].text = str(
                issue
            )


            row[1].text = str(
                count
            )


        document.add_paragraph()


        document.add_paragraph(
            (
                "Issue values represent detected "
                "issue mentions. One review may "
                "contain more than one issue."
            )
        )


    else:

        document.add_paragraph(
            (
                "No negative-review issue "
                "categories were identified."
            )
        )


    # ========================================================
    # 3. CUSTOMER VOICE
    # ========================================================

    document.add_heading(
        (
            "3. Customer Voice Intelligence"
        ),
        level=1,
    )


    # --------------------------------------------------------
    # POSITIVE TERMS
    # --------------------------------------------------------

    document.add_heading(
        (
            "3.1 Frequent Positive Terms"
        ),
        level=2,
    )


    positive_words = (
        analysis_summary.get(
            "top_positive_words",
            [],
        )
    )


    if positive_words:

        positive_table = (
            document.add_table(
                rows=1,
                cols=2,
            )
        )


        positive_table.style = (
            "Table Grid"
        )


        positive_table.alignment = (
            WD_TABLE_ALIGNMENT.CENTER
        )


        header = (
            positive_table.rows[
                0
            ].cells
        )


        header[0].text = (
            "Word"
        )


        header[1].text = (
            "Frequency"
        )


        for item in (
            positive_words
        ):

            (
                word,
                count,
            ) = get_word_and_count(
                item
            )


            row = (
                positive_table
                .add_row()
                .cells
            )


            row[0].text = str(
                word
            )


            row[1].text = str(
                count
            )


    else:

        document.add_paragraph(
            (
                "No positive customer-language "
                "terms were available."
            )
        )


    # --------------------------------------------------------
    # NEGATIVE TERMS
    # --------------------------------------------------------

    document.add_heading(
        (
            "3.2 Frequent Negative Terms"
        ),
        level=2,
    )


    negative_words = (
        analysis_summary.get(
            "top_negative_words",
            [],
        )
    )


    if negative_words:

        negative_table = (
            document.add_table(
                rows=1,
                cols=2,
            )
        )


        negative_table.style = (
            "Table Grid"
        )


        negative_table.alignment = (
            WD_TABLE_ALIGNMENT.CENTER
        )


        header = (
            negative_table.rows[
                0
            ].cells
        )


        header[0].text = (
            "Word"
        )


        header[1].text = (
            "Frequency"
        )


        for item in (
            negative_words
        ):

            (
                word,
                count,
            ) = get_word_and_count(
                item
            )


            row = (
                negative_table
                .add_row()
                .cells
            )


            row[0].text = str(
                word
            )


            row[1].text = str(
                count
            )


    else:

        document.add_paragraph(
            (
                "No negative customer-language "
                "terms were available."
            )
        )


    # ========================================================
    # 4. REPRESENTATIVE NEGATIVE REVIEWS
    # ========================================================

    document.add_heading(
        (
            "4. Representative Negative "
            "Customer Reviews"
        ),
        level=1,
    )


    negative_reviews = (
        analysis_summary.get(
            "sample_negative_reviews",
            [],
        )
    )


    if negative_reviews:

        for review in (
            negative_reviews
        ):

            document.add_paragraph(
                str(review),
                style="List Bullet",
            )


    else:

        document.add_paragraph(
            (
                "No representative negative "
                "reviews were available."
            )
        )


    # ========================================================
    # 5. DEPARTMENT REPORTS
    # ========================================================

    document.add_page_break()


    document.add_heading(
        (
            "5. AI Department Manager Reports"
        ),
        level=1,
    )


    document.add_paragraph(
        (
            "The department-level reports are "
            "generated through a multi-provider "
            "Large Language Model architecture. "
            "The provider and model used for each "
            "report are recorded for transparency."
        )
    )


    manager_order = [

        "Technical Manager",

        "Product Manager",

        "Customer Service Manager",

        "Marketing Manager",

        "Subscription Manager",
    ]


    for (
        number,
        manager_name,
    ) in enumerate(
        manager_order,
        start=1,
    ):

        report = (
            manager_reports.get(
                manager_name
            )
        )


        if not report:
            continue


        document.add_heading(
            (
                f"5.{number} "
                f"{manager_name}"
            ),
            level=2,
        )


        provider = (
            report.get(
                "provider",
                "Unknown",
            )
        )


        model = (
            report.get(
                "model",
                "Unknown",
            )
        )


        source = (
            document.add_paragraph()
        )


        provider_run = (
            source.add_run(
                "LLM Provider: "
            )
        )


        provider_run.bold = True


        source.add_run(
            str(provider)
        )


        source.add_run(
            " | "
        )


        model_run = (
            source.add_run(
                "Model: "
            )
        )


        model_run.bold = True


        source.add_run(
            str(model)
        )


        document.add_paragraph()


        docx_add_llm_report(
            document,
            report.get(
                "content",
                "",
            )
        )


        document.add_paragraph()


    # ========================================================
    # 6. EXECUTIVE REPORT
    # ========================================================

    document.add_page_break()


    document.add_heading(
        (
            "6. Executive Brand "
            "Reputation Report"
        ),
        level=1,
    )


    if executive_report:

        provider = (
            executive_report.get(
                "provider",
                "Unknown",
            )
        )


        model = (
            executive_report.get(
                "model",
                "Unknown",
            )
        )


        source = (
            document.add_paragraph()
        )


        provider_run = (
            source.add_run(
                "Executive LLM Provider: "
            )
        )


        provider_run.bold = True


        source.add_run(
            str(provider)
        )


        source.add_run(
            " | "
        )


        model_run = (
            source.add_run(
                "Model: "
            )
        )


        model_run.bold = True


        source.add_run(
            str(model)
        )


        document.add_paragraph()


        docx_add_llm_report(
            document,
            executive_report.get(
                "content",
                "",
            )
        )


    else:

        document.add_paragraph(
            (
                "The Executive Report "
                "has not been generated."
            )
        )


    # ========================================================
    # 7. SYSTEM INTERPRETATION
    # ========================================================

    document.add_page_break()


    document.add_heading(
        (
            "7. System Interpretation "
            "and Limitations"
        ),
        level=1,
    )


    interpretation_notes = [

        (
            "Sentiment classifications are "
            "predictions generated by the "
            "trained DistilBERT model."
        ),

        (
            "Model confidence represents a "
            "model output probability and does "
            "not guarantee prediction correctness."
        ),

        (
            "The Brand Reputation Score is a "
            "project-defined decision-support "
            "indicator and not a universal "
            "industry-standard metric."
        ),

        (
            "Issue categories are generated "
            "using predefined keyword-based "
            "issue-analysis logic."
        ),

        (
            "One negative review may contribute "
            "to multiple issue categories."
        ),

        (
            "Large Language Model recommendations "
            "are intended as decision-support "
            "outputs rather than verified business "
            "decisions."
        ),

        (
            "OpenRouter may route requests to "
            "different available free models."
        ),

        (
            "Ollama Cloud model availability and "
            "usage depend on the configured "
            "service limits."
        ),

        (
            "Gemini and other third-party AI "
            "services may experience temporary "
            "quota or rate-limit restrictions."
        ),
    ]


    for note in (
        interpretation_notes
    ):

        document.add_paragraph(
            note,
            style="List Bullet",
        )


    # ========================================================
    # SAVE DOCX TO MEMORY
    # ========================================================

    output = BytesIO()


    document.save(
        output
    )


    output.seek(
        0
    )


    return output


# ============================================================
# PDF CLASS
# ============================================================

class BrandPulsePDF(
    FPDF
):

    def header(
        self,
    ):
        """
        Display a header from page 2 onward.
        """

        if (
            self.page_no()
            <= 1
        ):

            return


        self.set_font(
            "Helvetica",
            "B",
            8.5,
        )


        self.set_text_color(
            30,
            120,
            85,
        )


        self.set_x(
            self.l_margin
        )


        self.cell(
            w=0,
            h=6,
            text=(
                "BrandPulse AI - "
                "Brand Reputation Intelligence Report"
            ),
            align="C",
            new_x="LMARGIN",
            new_y="NEXT",
        )


        self.set_draw_color(
            210,
            215,
            220,
        )


        self.line(
            self.l_margin,
            self.get_y(),
            self.w
            - self.r_margin,
            self.get_y(),
        )


        self.ln(
            3
        )


    def footer(
        self,
    ):
        """
        Display page number.
        """

        self.set_y(
            -14
        )


        self.set_font(
            "Helvetica",
            "",
            8,
        )


        self.set_text_color(
            110,
            110,
            110,
        )


        self.set_x(
            self.l_margin
        )


        self.cell(
            w=0,
            h=8,
            text=(
                f"Page {self.page_no()}"
            ),
            align="C",
        )


# ============================================================
# PDF POSITION HELPER
# ============================================================

def reset_pdf_x(
    pdf,
):
    """
    Always return the PDF cursor to the
    left page margin.

    This prevents:
    'Not enough horizontal space to render
    a single character'
    """

    pdf.set_x(
        pdf.l_margin
    )


# ============================================================
# PDF TEXT HELPERS
# ============================================================

def pdf_add_heading(
    pdf,
    text,
    level=1,
):

    text = pdf_safe_text(
        text
    )


    if not text:
        return


    if level == 1:

        pdf.ln(
            3
        )


        pdf.set_font(
            "Helvetica",
            "B",
            14,
        )


        pdf.set_text_color(
            30,
            100,
            80,
        )


        reset_pdf_x(
            pdf
        )


        pdf.multi_cell(
            w=0,
            h=7.5,
            text=text,
            new_x="LMARGIN",
            new_y="NEXT",
        )


        pdf.ln(
            2
        )


    elif level == 2:

        pdf.ln(
            2
        )


        pdf.set_font(
            "Helvetica",
            "B",
            11.5,
        )


        pdf.set_text_color(
            55,
            55,
            55,
        )


        reset_pdf_x(
            pdf
        )


        pdf.multi_cell(
            w=0,
            h=6.5,
            text=text,
            new_x="LMARGIN",
            new_y="NEXT",
        )


        pdf.ln(
            1
        )


    else:

        pdf.ln(
            1
        )


        pdf.set_font(
            "Helvetica",
            "B",
            10,
        )


        pdf.set_text_color(
            70,
            70,
            70,
        )


        reset_pdf_x(
            pdf
        )


        pdf.multi_cell(
            w=0,
            h=6,
            text=text,
            new_x="LMARGIN",
            new_y="NEXT",
        )


def pdf_add_paragraph(
    pdf,
    text,
):

    text = pdf_safe_text(
        text
    )


    if not text:
        return


    pdf.set_font(
        "Helvetica",
        "",
        9.3,
    )


    pdf.set_text_color(
        45,
        45,
        45,
    )


    reset_pdf_x(
        pdf
    )


    pdf.multi_cell(
        w=0,
        h=5.6,
        text=text,
        new_x="LMARGIN",
        new_y="NEXT",
    )


    pdf.ln(
        1.5
    )


def pdf_add_bullet(
    pdf,
    text,
):

    text = pdf_safe_text(
        text
    )


    if not text:
        return


    pdf.set_font(
        "Helvetica",
        "",
        9.2,
    )


    pdf.set_text_color(
        45,
        45,
        45,
    )


    reset_pdf_x(
        pdf
    )


    pdf.multi_cell(
        w=0,
        h=5.6,
        text=(
            "- "
            + text
        ),
        new_x="LMARGIN",
        new_y="NEXT",
    )


def pdf_add_llm_report(
    pdf,
    report_text,
):
    """
    Convert Markdown-style LLM reports
    into PDF content.
    """

    if not report_text:
        return


    lines = (
        str(report_text)
        .splitlines()
    )


    for raw_line in lines:

        line = (
            raw_line.strip()
        )


        if not line:
            continue


        # ----------------------------------------------------
        # H1
        # ----------------------------------------------------

        if line.startswith(
            "# "
        ):

            pdf_add_heading(
                pdf,
                line[2:],
                level=1,
            )


        # ----------------------------------------------------
        # H2
        # ----------------------------------------------------

        elif line.startswith(
            "## "
        ):

            pdf_add_heading(
                pdf,
                line[3:],
                level=2,
            )


        # ----------------------------------------------------
        # H3
        # ----------------------------------------------------

        elif line.startswith(
            "### "
        ):

            pdf_add_heading(
                pdf,
                line[4:],
                level=3,
            )


        # ----------------------------------------------------
        # BULLET
        # ----------------------------------------------------

        elif (
            line.startswith("- ")
            or
            line.startswith("• ")
        ):

            pdf_add_bullet(
                pdf,
                line[2:],
            )


        # ----------------------------------------------------
        # NUMBERED LIST
        # ----------------------------------------------------

        elif re.match(
            r"^\d+\.\s",
            line,
        ):

            pdf_add_paragraph(
                pdf,
                line,
            )


        # ----------------------------------------------------
        # NORMAL TEXT
        # ----------------------------------------------------

        else:

            pdf_add_paragraph(
                pdf,
                line,
            )


# ============================================================
# PDF TABLE
# ============================================================

def pdf_add_table(
    pdf,
    rows,
    column_widths=None,
):
    """
    Create a robust two-column table.

    Column widths are automatically scaled
    to the usable page width.
    """

    if not rows:
        return


    usable_width = (
        pdf.w
        - pdf.l_margin
        - pdf.r_margin
    )


    # ========================================================
    # COLUMN WIDTHS
    # ========================================================

    if (
        column_widths is None
        or
        len(
            column_widths
        ) != 2
    ):

        first_width = (
            usable_width
            * 0.70
        )

        second_width = (
            usable_width
            * 0.30
        )


    else:

        requested_total = sum(
            column_widths
        )


        if requested_total <= 0:

            first_width = (
                usable_width
                * 0.70
            )

            second_width = (
                usable_width
                * 0.30
            )


        else:

            scale_factor = (
                usable_width
                /
                requested_total
            )


            first_width = (
                column_widths[
                    0
                ]
                * scale_factor
            )


            second_width = (
                column_widths[
                    1
                ]
                * scale_factor
            )


    # Leave a small safety space.
    first_width -= (
        0.7
    )

    second_width -= (
        0.7
    )


    # ========================================================
    # TABLE ROWS
    # ========================================================

    for (
        row_number,
        row,
    ) in enumerate(
        rows
    ):

        reset_pdf_x(
            pdf
        )


        # ----------------------------------------------------
        # HEADER
        # ----------------------------------------------------

        if row_number == 0:

            pdf.set_fill_color(
                40,
                110,
                90,
            )


            pdf.set_text_color(
                255,
                255,
                255,
            )


            pdf.set_font(
                "Helvetica",
                "B",
                8.8,
            )


        # ----------------------------------------------------
        # BODY
        # ----------------------------------------------------

        else:

            pdf.set_fill_color(
                245,
                247,
                249,
            )


            pdf.set_text_color(
                45,
                45,
                45,
            )


            pdf.set_font(
                "Helvetica",
                "",
                8.8,
            )


        left_text = (
            pdf_safe_text(
                row[0]
            )
        )


        right_text = (
            pdf_safe_text(
                row[1]
            )
        )


        # Avoid very long continuous values.
        if len(
            left_text
        ) > 110:

            left_text = (
                left_text[
                    :107
                ]
                + "..."
            )


        if len(
            right_text
        ) > 65:

            right_text = (
                right_text[
                    :62
                ]
                + "..."
            )


        # ----------------------------------------------------
        # LEFT CELL
        # ----------------------------------------------------

        pdf.cell(
            w=first_width,
            h=7,
            text=left_text,
            border=1,
            fill=True,
        )


        # ----------------------------------------------------
        # RIGHT CELL
        # ----------------------------------------------------

        pdf.cell(
            w=second_width,
            h=7,
            text=right_text,
            border=1,
            fill=True,
            new_x="LMARGIN",
            new_y="NEXT",
        )


    pdf.ln(
        4
    )


# ============================================================
# CREATE PDF REPORT
# ============================================================

def create_brandpulse_pdf(
    analysis_summary,
    manager_reports,
    executive_report,
):
    """
    Create a complete BrandPulse AI PDF.

    Returns:
        bytes
    """

    pdf = BrandPulsePDF(
        orientation="P",
        unit="mm",
        format="A4",
    )


    # ========================================================
    # PAGE SETTINGS
    # ========================================================

    pdf.set_margins(
        left=15,
        top=15,
        right=15,
    )


    pdf.set_auto_page_break(
        auto=True,
        margin=18,
    )


    # ========================================================
    # COVER PAGE
    # ========================================================

    pdf.add_page()


    pdf.set_y(
        48
    )


    reset_pdf_x(
        pdf
    )


    pdf.set_font(
        "Helvetica",
        "B",
        27,
    )


    pdf.set_text_color(
        25,
        145,
        80,
    )


    pdf.cell(
        w=0,
        h=13,
        text="BrandPulse AI",
        align="C",
        new_x="LMARGIN",
        new_y="NEXT",
    )


    pdf.ln(
        8
    )


    pdf.set_font(
        "Helvetica",
        "B",
        14,
    )


    pdf.set_text_color(
        40,
        40,
        40,
    )


    reset_pdf_x(
        pdf
    )


    pdf.multi_cell(
        w=0,
        h=8,
        text=(
            "AI-Assisted Brand Reputation "
            "Intelligence Report"
        ),
        align="C",
        new_x="LMARGIN",
        new_y="NEXT",
    )


    pdf.ln(
        7
    )


    pdf.set_font(
        "Helvetica",
        "",
        11,
    )


    pdf.set_text_color(
        80,
        80,
        80,
    )


    reset_pdf_x(
        pdf
    )


    pdf.multi_cell(
        w=0,
        h=7,
        text=(
            "Online Review-Based Brand "
            "Reputation Prediction "
            "Using NLP Techniques"
        ),
        align="C",
        new_x="LMARGIN",
        new_y="NEXT",
    )


    pdf.ln(
        10
    )


    pdf.set_font(
        "Helvetica",
        "",
        9.5,
    )


    reset_pdf_x(
        pdf
    )


    pdf.cell(
        w=0,
        h=7,
        text=(
            "Predictive Model: DistilBERT"
        ),
        align="C",
        new_x="LMARGIN",
        new_y="NEXT",
    )


    reset_pdf_x(
        pdf
    )


    pdf.cell(
        w=0,
        h=7,
        text=(
            "Department AI: "
            "OpenRouter and Ollama"
        ),
        align="C",
        new_x="LMARGIN",
        new_y="NEXT",
    )


    reset_pdf_x(
        pdf
    )


    pdf.cell(
        w=0,
        h=7,
        text=(
            "Executive AI: Gemini"
        ),
        align="C",
        new_x="LMARGIN",
        new_y="NEXT",
    )


    pdf.ln(
        6
    )


    pdf.set_text_color(
        110,
        110,
        110,
    )


    reset_pdf_x(
        pdf
    )


    pdf.cell(
        w=0,
        h=7,
        text=(
            "Generated: "
            + current_timestamp()
        ),
        align="C",
        new_x="LMARGIN",
        new_y="NEXT",
    )


    # ========================================================
    # PAGE 2
    # ========================================================

    pdf.add_page()


    # ========================================================
    # 1. BRAND REPUTATION OVERVIEW
    # ========================================================

    pdf_add_heading(
        pdf,
        (
            "1. Brand Reputation Overview"
        ),
        level=1,
    )


    overview_rows = [

        [
            "Indicator",
            "Result",
        ],

        [
            "Reviews Analysed",
            str(
                analysis_summary.get(
                    "total_reviews",
                    0,
                )
            ),
        ],

        [
            "Positive Reviews",
            str(
                analysis_summary.get(
                    "positive_reviews",
                    0,
                )
            ),
        ],

        [
            "Negative Reviews",
            str(
                analysis_summary.get(
                    "negative_reviews",
                    0,
                )
            ),
        ],

        [
            "Positive Percentage",
            (
                f"{analysis_summary.get(
                    'positive_percentage',
                    0
                )}%"
            ),
        ],

        [
            "Negative Percentage",
            (
                f"{analysis_summary.get(
                    'negative_percentage',
                    0
                )}%"
            ),
        ],

        [
            "Brand Reputation Score",
            (
                f"{analysis_summary.get(
                    'reputation_score',
                    0
                )}%"
            ),
        ],
    ]


    pdf_add_table(
        pdf,
        overview_rows,
        [
            115,
            65,
        ],
    )


    pdf_add_paragraph(
        pdf,
        (
            "Methodology Note: The Brand "
            "Reputation Score is a project-defined "
            "decision-support indicator calculated "
            "from the proportion of positive "
            "DistilBERT sentiment predictions. "
            "It is not an official Spotify or "
            "industry-standard reputation metric."
        ),
    )


    # ========================================================
    # 2. ISSUE ANALYSIS
    # ========================================================

    pdf_add_heading(
        pdf,
        (
            "2. Negative Review Issue Analysis"
        ),
        level=1,
    )


    issue_counts = (
        analysis_summary.get(
            "issue_counts",
            {},
        )
    )


    if issue_counts:

        issue_rows = [

            [
                "Issue Category",
                "Mentions",
            ]
        ]


        for (
            issue,
            count,
        ) in issue_counts.items():

            issue_rows.append(
                [
                    str(
                        issue
                    ),
                    str(
                        count
                    ),
                ]
            )


        pdf_add_table(
            pdf,
            issue_rows,
            [
                135,
                45,
            ],
        )


        pdf_add_paragraph(
            pdf,
            (
                "Issue values represent detected "
                "issue mentions. A single review "
                "may contain more than one issue."
            ),
        )


    else:

        pdf_add_paragraph(
            pdf,
            (
                "No negative-review issue "
                "categories were identified."
            ),
        )


    # ========================================================
    # 3. CUSTOMER VOICE
    # ========================================================

    pdf_add_heading(
        pdf,
        (
            "3. Customer Voice Intelligence"
        ),
        level=1,
    )


    # --------------------------------------------------------
    # POSITIVE
    # --------------------------------------------------------

    pdf_add_heading(
        pdf,
        (
            "3.1 Frequent Positive Terms"
        ),
        level=2,
    )


    positive_words = (
        analysis_summary.get(
            "top_positive_words",
            [],
        )
    )


    if positive_words:

        rows = [

            [
                "Word",
                "Frequency",
            ]
        ]


        for item in (
            positive_words
        ):

            (
                word,
                count,
            ) = get_word_and_count(
                item
            )


            rows.append(
                [
                    str(
                        word
                    ),
                    str(
                        count
                    ),
                ]
            )


        pdf_add_table(
            pdf,
            rows,
            [
                135,
                45,
            ],
        )


    else:

        pdf_add_paragraph(
            pdf,
            (
                "No positive customer-language "
                "terms were available."
            ),
        )


    # --------------------------------------------------------
    # NEGATIVE
    # --------------------------------------------------------

    pdf_add_heading(
        pdf,
        (
            "3.2 Frequent Negative Terms"
        ),
        level=2,
    )


    negative_words = (
        analysis_summary.get(
            "top_negative_words",
            [],
        )
    )


    if negative_words:

        rows = [

            [
                "Word",
                "Frequency",
            ]
        ]


        for item in (
            negative_words
        ):

            (
                word,
                count,
            ) = get_word_and_count(
                item
            )


            rows.append(
                [
                    str(
                        word
                    ),
                    str(
                        count
                    ),
                ]
            )


        pdf_add_table(
            pdf,
            rows,
            [
                135,
                45,
            ],
        )


    else:

        pdf_add_paragraph(
            pdf,
            (
                "No negative customer-language "
                "terms were available."
            ),
        )


    # ========================================================
    # 4. REPRESENTATIVE REVIEWS
    # ========================================================

    pdf_add_heading(
        pdf,
        (
            "4. Representative Negative "
            "Customer Reviews"
        ),
        level=1,
    )


    negative_reviews = (
        analysis_summary.get(
            "sample_negative_reviews",
            [],
        )
    )


    if negative_reviews:

        for (
            number,
            review,
        ) in enumerate(
            negative_reviews,
            start=1,
        ):

            pdf_add_paragraph(
                pdf,
                (
                    f"{number}. "
                    f"{review}"
                ),
            )


    else:

        pdf_add_paragraph(
            pdf,
            (
                "No representative negative "
                "reviews were available."
            ),
        )


    # ========================================================
    # 5. DEPARTMENT REPORTS
    # ========================================================

    pdf.add_page()


    pdf_add_heading(
        pdf,
        (
            "5. AI Department Manager Reports"
        ),
        level=1,
    )


    pdf_add_paragraph(
        pdf,
        (
            "Department reports are generated "
            "through a multi-provider LLM "
            "architecture. The provider and "
            "actual model used are recorded "
            "for transparency."
        ),
    )


    manager_order = [

        "Technical Manager",

        "Product Manager",

        "Customer Service Manager",

        "Marketing Manager",

        "Subscription Manager",
    ]


    for (
        manager_number,
        manager_name,
    ) in enumerate(
        manager_order,
        start=1,
    ):

        report = (
            manager_reports.get(
                manager_name
            )
        )


        if not report:
            continue


        # Start a new page if the current
        # one is nearly full.
        if (
            pdf.get_y()
            > 225
        ):

            pdf.add_page()


        pdf_add_heading(
            pdf,
            (
                f"5.{manager_number} "
                f"{manager_name}"
            ),
            level=2,
        )


        provider = (
            report.get(
                "provider",
                "Unknown",
            )
        )


        model = (
            report.get(
                "model",
                "Unknown",
            )
        )


        pdf_add_paragraph(
            pdf,
            (
                f"LLM Provider: {provider} "
                f"| Model: {model}"
            ),
        )


        pdf_add_llm_report(
            pdf,
            report.get(
                "content",
                "",
            ),
        )


        pdf.ln(
            3
        )


    # ========================================================
    # 6. EXECUTIVE REPORT
    # ========================================================

    pdf.add_page()


    pdf_add_heading(
        pdf,
        (
            "6. Executive Brand "
            "Reputation Report"
        ),
        level=1,
    )


    if executive_report:

        provider = (
            executive_report.get(
                "provider",
                "Unknown",
            )
        )


        model = (
            executive_report.get(
                "model",
                "Unknown",
            )
        )


        pdf_add_paragraph(
            pdf,
            (
                "Executive LLM Provider: "
                f"{provider} "
                f"| Model: {model}"
            ),
        )


        pdf_add_llm_report(
            pdf,
            executive_report.get(
                "content",
                "",
            ),
        )


    else:

        pdf_add_paragraph(
            pdf,
            (
                "The Executive Report "
                "has not been generated."
            ),
        )


    # ========================================================
    # 7. INTERPRETATION / LIMITATIONS
    # ========================================================

    pdf.add_page()


    pdf_add_heading(
        pdf,
        (
            "7. System Interpretation "
            "and Limitations"
        ),
        level=1,
    )


    limitation_notes = [

        (
            "Sentiment classifications are "
            "predictions produced by the "
            "trained DistilBERT model."
        ),

        (
            "Model confidence represents the "
            "model output probability and does "
            "not guarantee prediction correctness."
        ),

        (
            "The Brand Reputation Score is "
            "a project-defined decision-support "
            "indicator rather than a universal "
            "industry-standard metric."
        ),

        (
            "Issue categories are generated "
            "using predefined keyword-based "
            "issue-analysis logic."
        ),

        (
            "A single negative review may "
            "contribute to multiple issue "
            "categories."
        ),

        (
            "LLM-generated recommendations "
            "are intended as decision-support "
            "outputs rather than verified "
            "management decisions."
        ),

        (
            "OpenRouter may select different "
            "free models when the openrouter/free "
            "route is requested."
        ),

        (
            "Ollama Cloud model availability "
            "and usage depend on configured "
            "service limitations."
        ),

        (
            "Gemini and other third-party "
            "LLM providers may experience "
            "temporary quota or rate-limit "
            "restrictions."
        ),
    ]


    for note in (
        limitation_notes
    ):

        pdf_add_bullet(
            pdf,
            note,
        )


    # ========================================================
    # FINAL ACADEMIC NOTE
    # ========================================================

    pdf.ln(
        7
    )


    pdf_add_paragraph(
        pdf,
        (
            "BrandPulse AI is an academic "
            "Final Year Project prototype "
            "designed to demonstrate the "
            "integration of predictive NLP, "
            "brand-reputation analytics and "
            "generative AI decision support."
        ),
    )


    # ========================================================
    # OUTPUT
    # ========================================================

    result = (
        pdf.output()
    )


    if isinstance(
        result,
        bytearray,
    ):

        result = bytes(
            result
        )


    return result
